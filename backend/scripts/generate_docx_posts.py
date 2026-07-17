# backend/scripts/generate_docx_posts.py
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_posts_document():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 1. Document Title / Header
    title = doc.add_paragraph()
    title_run = title.add_run("RafApp — LinkedIn Marketing Content Kit")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)  # Indigo
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("10 Engaging Posts for Promoting the Industrial OS for Electrical Contractors in Iceland")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)  # Gray
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("").paragraph_format.space_after = Pt(24)

    # Paths to generated graphics
    img1 = r"C:\Users\mario\.gemini\antigravity\brain\a4a1acf0-918f-45cf-9d9e-a218510c0f47\rafapp_linkedin_promo_1_1784151702244.png"
    img2 = r"C:\Users\mario\.gemini\antigravity\brain\a4a1acf0-918f-45cf-9d9e-a218510c0f47\rafapp_linkedin_promo_2_1784151712784.png"
    img3 = r"C:\Users\mario\.gemini\antigravity\brain\a4a1acf0-918f-45cf-9d9e-a218510c0f47\rafapp_linkedin_promo_3_1784151725305.png"

    posts = [
        {
            "num": 1,
            "title": "RafApp: Framtíðin í rekstri rafverktaka á Íslandi ⚡",
            "title_en": "RafApp: The Future of Electrical Contracting in Iceland ⚡",
            "icelandic": "Er fyrirtækið þitt enn að treysta á gamaldags Excel-skjöl, pappírsvinnu og dreifð gögn?\n\nKynntu þér RafApp – eina heildstæða stýrikerfið sem er hannað sérstaklega fyrir íslenska rafverktaka. Frá verkefnastýringu og efnisflæði yfir í flókna ákvæðisvinnu rafvirkja, RafApp sameinar allt á einum stað.\n\nSparaðu tíma, fækkaðu mistökum á skrifstofunni og auka framleiðni úti á verkstað.",
            "english": "Is your business still relying on outdated spreadsheets, paperwork, and scattered data?\n\nMeet RafApp – the ultimate OS designed specifically for Icelandic electrical contractors. From project management and materials to complex piecework settlements, RafApp brings everything together. Save time, reduce errors, and increase efficiency on-site.",
            "hashtags": "#Rafverktakar #RafApp #Verkefnastýring #Framkvæmdir #SaaS #Iceland #TechnicianOS",
            "cta": "Skoðaðu nánar á www.rafapp.is og bókaðu kynningu strax í dag!",
            "image": img1,
            "image_caption": "Mynd 1: RafApp nútímalegt mælaborð og stjórnunarkerfi."
        },
        {
            "num": 2,
            "title": "Sjálfvirk ákvæðisvinna: Einn smellur í stað margra klukkustunda handavinnu 📊",
            "title_en": "Automated Piecework: One Click Instead of Hours of Spreadsheets 📊",
            "icelandic": "Ákvæðisvinna rafvirkja (samkvæmt samningum RSÍ og SART) er flókin og tímafrek reikniformúla. Hæðarálög, steypuálög, reiknitala... handvirkir útreikningar geta tekið marga tíma og boðið upp á dýr mistök.\n\nRafApp leysir þetta með sjálfvirkri reiknivél sem reiknar út uppgjör samstundis byggt á skráðum verkum. Það gefur bæði vinnuflokknum og verktakanum fullkomið gagnsæi og rétt uppgjör á svipstundu.",
            "english": "Electrician piecework calculations under RSÍ and SART regulations are notoriously complex. Surcharges for heights, concrete grids, and calculations can take hours. RafApp automates this process instantly based on completed logs, ensuring total accuracy and transparency for both crew and contractor.",
            "hashtags": "#Ákvæðisvinna #RSÍ #SART #Rafvirkjar #Verktakar #Sjálfvirkni #Launauppgjör",
            "cta": "Láttu reiknivélina vinna vinnuna fyrir þig. Kynntu þér málið á www.rafapp.is.",
            "image": img3,
            "image_caption": "Mynd 2: Sjálfvirk ákvæðisvinnureiknivél í RafApp."
        },
        {
            "num": 3,
            "title": "Rafvirkjar með stjórnina í símanum á verkstað 📱",
            "title_en": "Electricians in Control: Mobile App for Jobsites 📱",
            "icelandic": "Með farsímaforriti RafApp geta rafvirkjar á verkstað auðveldlega skráð lokið verk, skráð vinnutíma og fylgst með efnisnotkun í rauntíma.\n\nEngir tapaðir seðlar eða óskýrar handskrifaðar glósur. Allar upplýsingar fara beint inn í kerfið svo skrifstofan hefur fullkomna yfirsýn og getur undirbúið reikninga strax.",
            "english": "With RafApp's mobile application, electricians on-site can easily log completed tasks, track work hours, and monitor material usage in real-time. No lost papers, no handwriting deciphering. Everything flows straight to the office for immediate invoicing.",
            "hashtags": "#Snjalllausnir #Farsímaforrit #Verkstaður #RafApp #Vinnuskráning #MobileWorkforce",
            "cta": "Bættu samskiptin milli verkstaðar og skrifstofu á www.rafapp.is.",
            "image": img2,
            "image_caption": "Mynd 3: Farsímaforrit RafApp á vettvangi."
        },
        {
            "num": 4,
            "title": "Hvar er efnið? Fullkomin efnisskráning og birgðastýring 📦",
            "title_en": "Where's the Material? Seamless Material Tracking 📦",
            "icelandic": "Tapast efni á verkstigi eða gleymist að rukka það? Það kostar verktaka mikla peninga árlega.\n\nRafApp heldur utan um birgðir og verkefnatengt efnisflæði í rauntíma. Þú sérð nákvæmlega hvaða efni fór á hvaða verkefni, hvað er til á lager og hvenær þarf að panta meira. Engin sóun, aðeins hámarksnýtni.",
            "english": "Losing materials or forgetting to bill them? It costs contractors money every year. RafApp keeps track of inventory and project-specific material flow in real-time. Know exactly what went where, what's in stock, and when to reorder.",
            "hashtags": "#Birgðastýring #Efnisskráning #RafApp #Framkvæmdir #Nýtni #Logistics",
            "cta": "Náðu stjórn á efniskostnaðinum með RafApp. Skoðaðu nánar á www.rafapp.is.",
            "image": None
        },
        {
            "num": 5,
            "title": "Bílaleigubækur og dekkjastýring í einu kerfi 🚗",
            "title_en": "Car Fleet & Tyre Management in One OS 🚗",
            "icelandic": "Ertu með marga bíla í gangi á mismunandi verkstöðum? Að halda utan um viðhald getur verið flókið.\n\nRafApp hjálpar þér að halda utan um bílaflotann, akstursbækur, smurþjónustu og dekkjaskipti. Starfsmenn skrá kílómetrastöðu og smurbækur beint í appið svo flotinn sé ávallt í toppstandi.",
            "english": "Managing multiple vehicles across jobsites? Keeping track of maintenance is tough. RafApp helps you monitor your car fleet, mileage, oil changes, and tyre status. Employees log details directly in the app so your fleet is always road-ready.",
            "hashtags": "#Bílafloti #Dekkjastýring #Flotastýring #Rafverktakar #Stafrænt #FleetManagement",
            "cta": "Prófaðu stafræna flotastýringu á www.rafapp.is.",
            "image": None
        },
        {
            "num": 6,
            "title": "Skilvirk tilboðagerð án fyrirhafnar 💼",
            "title_en": "Effortless B2B Offer Generation 💼",
            "icelandic": "Þarftu að senda faglegt tilboð fljótt? Handvirk skrif tefja ferlið.\n\nRafApp gerir þér kleift að búa til tilboð beint úr efnis- og vinnuliðaskrá, á íslensku eða ensku. Kerfið reiknar nákvæmlega einingar og verð út frá nýjustu einingatöflum svo þú getir sent tilboð með örfáum smellum.",
            "english": "Need to send professional proposals quickly? Manual drafting delays sales. RafApp allows you to generate offers directly from your labor and material catalogs, in both Icelandic and English. The system calculates exact units and rates so you can win bids faster.",
            "hashtags": "#Tilboðagerð #Verkefni #Rafverktakar #B2B #Skilvirkni #SalesAutomation",
            "cta": "Hafðu hraðann með þér í tilboðagerðinni. Sjá nánar á www.rafapp.is.",
            "image": None
        },
        {
            "num": 7,
            "title": "Einfaldari launakeyrsla og rétt uppgjör 💳",
            "title_en": "Simpler Payroll and Accurate Settlements 💳",
            "icelandic": "Ónákvæmar tímaskráningar og flóknir bónusar valda oft ágreiningi við launakeyrslu.\n\nMeð RafApp er hver vinnustund, ferðatími og ákvæðisvinnubónus skráður á rétt verkefni. Kerfið dregur vinnulaunafyrirframgreiðslur sjálfkrafa frá ákvæðisvinnuuppgjörinu, sem tryggir réttan launaseðil í hvert einasta skipti.",
            "english": "Inaccurate time logs and complex bonuses often lead to payroll disputes. With RafApp, every work hour, travel time, and piecework bonus is logged against the correct project. The system automatically handles wage advance deductions, ensuring accurate payslips every time.",
            "hashtags": "#Tímaskráning #Launaseðlar #Launakeyrsla #RafApp #Launamál #FinTech",
            "cta": "Launakeyrsla án áhyggna á www.rafapp.is.",
            "image": None
        },
        {
            "num": 8,
            "title": "Öryggi gagna í fyrirrúmi: 2FA og örugg skýjalausn 🔒",
            "title_en": "Data Security First: 2FA Cloud Infrastructure 🔒",
            "icelandic": "Gögnin um verkin þín, viðskiptavini, teikningar og starfsmenn eru mikilvægustu eignir fyrirtækisins.\n\nRafApp verndar þessi gögn með tveggja-þátta auðkenningu (2FA) og öruggri dulkóðun í skýinu. Hýst á öruggum netþjónum svo þú hafir ávallt aðgang en óviðkomandi séu útilokaðir.",
            "english": "Your project data, drawings, client information, and employee files are valuable assets. RafApp protects them with two-factor authentication (2FA) and enterprise-grade cloud encryption. Hosted securely so you have access anytime, while keeping unauthorized access locked out.",
            "hashtags": "#Gagnaöryggi #2FA #Skýjalausnir #Netöryggi #RafApp #CloudSecurity",
            "cta": "Treystu á öruggan rekstur. Skráðu þig á www.rafapp.is.",
            "image": None
        },
        {
            "num": 9,
            "title": "Segðu bless við pappírsflóðið á verkstað 📄❌",
            "title_en": "Say Goodbye to Paper Blueprints and Loose Notes 📄❌",
            "icelandic": "Glósur á servíettum, ljósmyndir í einkasímum og úreltar pappírsteikningar hægja á verkinu.\n\nMeð RafApp skráir starfsfólkið framvindu með myndum, athugasemdum og gátlistum beint í forritið undir réttu verki. Allt er vistað á einum stað og aðgengilegt fyrir alla sem að verkinu koma.",
            "english": "Napkin notes, personal phone photos, and outdated paper prints slow down execution. With RafApp, teams capture progress with photos, comments, and checklists directly in the app. Everything is stored under the correct project folder for instant access.",
            "hashtags": "#Pappírslaust #Framkvæmdir #Tæknilausnir #RafApp #Gæðamál #DigitalTransformation",
            "cta": "Stafræna verkstaðinn með okkur á www.rafapp.is.",
            "image": None
        },
        {
            "num": 10,
            "title": "Hver er raunveruleg arðsemi af því að nota RafApp? 📈",
            "title_en": "What is the Real ROI of Using RafApp? 📈",
            "icelandic": "Verktakar sem nota RafApp spara að meðaltali 15-20 klukkustundir á mánuði í stjórnunarvinnu per starfsmann.\n\nMinni sóun á efni, nákvæmari tilboð, fljótari launakeyrsla og sjálfvirkt ákvæðisvinnuuppgjör þýða beinan sparnað í rekstri og bætta framlegð frá fyrsta degi.",
            "english": "Contractors using RafApp save an average of 15-20 hours a month in administrative tasks per employee. Less material waste, more accurate proposals, faster payroll, and automated piecework calculation translate directly into improved bottom-line margins.",
            "hashtags": "#ROI #Arðsemi #Sjálfvirkni #Rekstur #Rafverktakar #RafApp #BusinessOptimization",
            "cta": "Reiknaðu út þinn sparnað með því að heimsækja www.rafapp.is.",
            "image": None
        }
    ]

    for post in posts:
        # Add heading
        h = doc.add_heading(level=2)
        h_run = h.add_run(f"Post #{post['num']}: {post['title']}")
        h_run.font.name = 'Calibri'
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)

        # English translation of heading
        h_sub = doc.add_paragraph()
        h_sub_run = h_sub.add_run(f"English: {post['title_en']}")
        h_sub_run.font.name = 'Calibri'
        h_sub_run.font.size = Pt(10)
        h_sub_run.font.italic = True
        h_sub_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
        h_sub.paragraph_format.space_after = Pt(12)

        # Icelandic copy
        p_is_lbl = doc.add_paragraph()
        r = p_is_lbl.add_run("[ICELANDIC COPY]")
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        p_is_lbl.paragraph_format.space_after = Pt(2)

        p_is = doc.add_paragraph()
        p_is.add_run(post['icelandic'])
        p_is.paragraph_format.space_after = Pt(12)

        # English copy
        p_en_lbl = doc.add_paragraph()
        r = p_en_lbl.add_run("[ENGLISH COPY]")
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)
        p_en_lbl.paragraph_format.space_after = Pt(2)

        p_en = doc.add_paragraph()
        p_en_run = p_en.add_run(post['english'])
        p_en_run.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)
        p_en.paragraph_format.space_after = Pt(12)

        # Hashtags & CTA
        p_tags = doc.add_paragraph()
        r_tags = p_tags.add_run(f"Hashtags: {post['hashtags']}\n")
        r_tags.font.color.rgb = RGBColor(0x0d, 0x94, 0x88) # Teal
        r_tags.font.bold = True

        r_cta = p_tags.add_run(f"Call-to-Action: {post['cta']}")
        r_cta.font.bold = True
        r_cta.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
        p_tags.paragraph_format.space_after = Pt(12)

        # Insert image if any
        if post['image'] and os.path.exists(post['image']):
            try:
                # Add spacing
                doc.add_paragraph("").paragraph_format.space_after = Pt(6)
                # Add picture
                doc.add_picture(post['image'], width=Inches(5.5))
                # Add caption
                caption = doc.add_paragraph()
                caption_run = caption.add_run(post['image_caption'])
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(18)
            except Exception as e:
                print(f"Could not load image {post['image']}: {e}")

        # Page separator
        doc.add_page_break()

    # Save to Desktop
    desktop_path = r"c:\Users\mario\Desktop"
    output_filename = os.path.join(desktop_path, "RafApp_LinkedIn_Posts.docx")
    doc.save(output_filename)
    print(f"Word document saved to: {output_filename}")

if __name__ == "__main__":
    create_posts_document()
