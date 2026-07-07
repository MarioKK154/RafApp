import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_professional_guide():
    doc = docx.Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style helper
    def set_font(run, name="Segoe UI", size=11, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(10)
    run = title_p.add_run("RAFAPP FEATURES & CAPABILITIES GUIDE")
    set_font(run, size=24, bold=True, color=RGBColor(30, 41, 59)) # Dark Slate

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(200)
    run = subtitle_p.add_run("Comprehensive Product Document for Potential Customers\nProfessional Features, Workflows, and System Modules Overview")
    set_font(run, size=13, italic=True, color=RGBColor(107, 114, 128))

    doc.add_page_break()

    # Introduction
    intro_h = doc.add_paragraph()
    intro_h.paragraph_format.space_before = Pt(12)
    intro_h.paragraph_format.space_after = Pt(8)
    run = intro_h.add_run("Executive System Overview")
    set_font(run, size=16, bold=True, color=RGBColor(30, 41, 59))

    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(18)
    run = intro_p.add_run(
        "RafApp is an enterprise-grade Software-as-a-Service (SaaS) platform tailored specifically "
        "for electrical contracting companies, engineering firms, and utility managers. The platform "
        "bridges the gap between construction sites and back-office operations by consolidating "
        "project tracking, cost estimation, worker scheduling, fleet management, tool registries, "
        "procurement pipelines, and human resources in a unified, modern interface. This document "
        "provides a detailed look at every core functional module and user action available in the system."
    )
    set_font(run, size=11, color=RGBColor(55, 65, 81))

    # Modules
    modules = [
        {
            "num": "1",
            "title": "Tenant Architecture & Feature Toggles",
            "desc": "The platform is built on a multi-tenant database schema, allowing isolated enterprise accounts to run on shared infrastructure safely.",
            "actions": [
                {
                    "name": "Feature Flags Selection",
                    "detail": "Superadministrators can enable or disable functional modules (e.g. 'Fleet Management', 'Tool Registry') on a per-tenant basis. Changes update immediately, dynamically hiding or showing sidebar tabs for that tenant's users."
                },
                {
                    "name": "Strict Tenant Separation",
                    "detail": "Regular users can only access endpoints and data mapped to their specific tenant ID. Login access is strictly checked to prevent cross-tenant data leaks."
                }
            ]
        },
        {
            "num": "2",
            "title": "Project Management & Gantt Timelines",
            "desc": "Organize work sites, construct timelines, assign teams, and monitor project milestones from a central registry.",
            "actions": [
                {
                    "name": "Project Creation",
                    "detail": "Define project parameters including Name, Client, Physical Address, Start Date, Target End Date, and designated Project Manager."
                },
                {
                    "name": "Team Allocations",
                    "detail": "Assign specific engineers, leaders, and electricians to the project. Only assigned members can log work hours, view project details, or submit material requests."
                },
                {
                    "name": "Gantt Timeline Chart",
                    "detail": "Displays a visual project timeline with tasks, phases, and milestones. Gantt charts map out dependencies and help managers spot scheduling conflicts or resource constraints instantly."
                }
            ]
        },
        {
            "num": "3",
            "title": "Task Boards & Progress Tracking",
            "desc": "Create operational work items under projects, assign team members, and track status flow using Kanban-style boards.",
            "actions": [
                {
                    "name": "Task Creation and Assignment",
                    "detail": "Create individual tasks (e.g. 'Install Cable Trays in Block B'), set priority levels (Low, Medium, High), and assign them to electricians."
                },
                {
                    "name": "Kanban Status Transitions",
                    "detail": "Move tasks through lifecycle steps: To Do, In Progress, In Review, and Done. Electricians can update task status in real time from the field."
                },
                {
                    "name": "Photo Attachments & Uploads",
                    "detail": "Attach progress photos or damage reports directly to tasks. Managers can inspect the physical work completed without visiting the construction site."
                }
            ]
        },
        {
            "num": "4",
            "title": "Estimates, Pricing, and Labor Surcharges (Álagshlutföll)",
            "desc": "Prepare detailed commercial offers using automated labor rate calculations and catalog materials.",
            "actions": [
                {
                    "name": "Materials and Labor Entry",
                    "detail": "Search the unified catalog to add items (cables, pipes, fittings) and assign labor time codes (standardized minutes for installation per unit) directly to the proposal."
                },
                {
                    "name": "Work Load Modifiers (Álagshlutföll)",
                    "detail": "Toggle project modifier checkboxes in the sidebar (e.g., Overtime, Night Shifts, High-Altitude work, or Hazardous Conditions) to automatically scale labor costs and apply statutory price factors."
                },
                {
                    "name": "PDF Export and Printing",
                    "detail": "Compile estimate items into a professional, print-ready PDF commercial proposal with itemized costs, tax calculations, and term agreements for the customer."
                }
            ]
        },
        {
            "num": "5",
            "title": "Roster Scheduling & Calendar Event Organizer",
            "desc": "Manage the workforce roster and schedule company events, safety meetings, and site visits.",
            "actions": [
                {
                    "name": "Roster Planning",
                    "detail": "Construct weekly and monthly work schedules for electricians, mapping teams to specific projects to prevent over-allocation."
                },
                {
                    "name": "Past Date Booking Restrictions",
                    "detail": "The system prevents retroactive calendar scheduling. Meeting invites and resource bookings must be made for future dates only, protecting historical audit trails."
                },
                {
                    "name": "Event & Meeting Invites",
                    "detail": "Schedule meetings, training sessions, or customer site checkups. Attendees receive notifications on their dashboards and mobile apps."
                }
            ]
        },
        {
            "num": "6",
            "title": "Timesheets & Labor Approval Workflow",
            "desc": "Collect site check-in times and manual logs from electricians, verifying them for payroll processing.",
            "actions": [
                {
                    "name": "Mobile Work Logging",
                    "detail": "Electricians log hours worked per project and task, noting progress. Entries can be logged live or submitted manually at the end of the shift."
                },
                {
                    "name": "Manager Review & Approval",
                    "detail": "Managers review hours on their approval panel, verifying the logs against project milestones before approving them."
                },
                {
                    "name": "Timesheet Reporting & Export",
                    "detail": "Generate aggregated timesheet files for accounting. Approved hours flow directly into the invoicing and payroll systems."
                }
            ]
        },
        {
            "num": "7",
            "title": "HR, Payroll, & Absences",
            "desc": "Manage employee personnel files, sick leaves, paid time off, and monthly payslip generation.",
            "actions": [
                {
                    "name": "Absence Requests & Approvals",
                    "detail": "Electricians submit requests for holidays, sick leave, or parental leave. Managers review and approve them, automatically marking the user unavailable on scheduling boards."
                },
                {
                    "name": "Payslip Generation",
                    "detail": "Computes monthly base pay, overtime surcharges, and tax deductions from approved timesheets to generate formal payslips."
                }
            ]
        },
        {
            "num": "8",
            "title": "Shared Tools Registry & Logistics",
            "desc": "A central database of all capital tools, mapping custodianship and maintenance records.",
            "actions": [
                {
                    "name": "Tool Registration",
                    "detail": "Register power tools, testers, and crimpers with Serial Numbers, Purchase Date, and status flags (Ready, In Service, Damaged)."
                },
                {
                    "name": "Custodianship Tracking (Check-Out / Check-In)",
                    "detail": "Log check-outs when an electrician takes a tool. The system records who has the tool and when it is expected back, ensuring accountability."
                },
                {
                    "name": "Maintenance & Fault Reports",
                    "detail": "Users can flag damaged tools in the field and upload photos of broken parts. The system automatically notifies the workshop manager to arrange repairs."
                }
            ]
        },
        {
            "num": "9",
            "title": "Fleet Management",
            "desc": "Manage company vehicles, track mileage usage, tyre types, and report body damage.",
            "actions": [
                {
                    "name": "Mileage Book (Akstursbók)",
                    "detail": "Drivers enter starting and ending odometer values. The system logs distances traveled, updating maintenance reminders (e.g. oil change alerts) automatically."
                },
                {
                    "name": "Seasonal Tyre Tracker",
                    "detail": "Tracks whether a vehicle is currently fitted with summer tyres or studded winter tyres, ensuring company-wide compliance with winter driving laws."
                },
                {
                    "name": "Vehicle Incident Reports",
                    "detail": "Submit vehicle damage logs (like windshield cracks or fender benders) with photos directly from the mobile interface."
                }
            ]
        },
        {
            "num": "10",
            "title": "Procurement & Direct Supplier Shop Integrations",
            "desc": "Streamline material sourcing from construction sites with direct links to major electrical wholesalers.",
            "actions": [
                {
                    "name": "Material Requests Creation",
                    "detail": "Electricians on site create a procurement draft list, choosing items directly from the catalog database."
                },
                {
                    "name": "Direct Shop Linking",
                    "detail": "Catalog products include direct hyperlinks to partner supplier sites (Ískraft, Reykjafell, and Johan Ronning), enabling instant stock availability, spec, and pricing checks."
                },
                {
                    "name": "Purchase Order Approvals",
                    "detail": "Managers review request drafts, sign off, and submit order lists directly to suppliers for delivery dispatch."
                }
            ]
        },
        {
            "num": "11",
            "title": "Audit Logs, Compliance, & Superuser Impersonation",
            "desc": "Ensure enterprise-grade security and full compliance with strict logs and tracking.",
            "actions": [
                {
                    "name": "Superuser Impersonation Alerts",
                    "detail": "When a superuser impersonates a regular user's session to assist with a support issue, the system logs the action and automatically posts a notification to the target user's feed for auditing."
                },
                {
                    "name": "Session Audit Trail",
                    "detail": "Every data creation, deletion, or critical setting update is timestamped and logged, providing an unalterable audit trail for audit checks."
                }
            ]
        },
        {
            "num": "12",
            "title": "Financial Systems & Direct Payment Operations",
            "desc": "Managing software subscription billing, company transactions, and payroll distribution workflows.",
            "actions": [
                {
                    "name": "SaaS Subscription Payments",
                    "detail": "Companies pay for RafApp subscriptions directly through the app using integrated card payment processors (Stripe or Rapyd). Invoices and receipts are generated automatically."
                },
                {
                    "name": "Direct Financial Transactions",
                    "detail": "RafApp acts as an operational management system, not a banking gateway. Companies cannot make direct financial transactions with other companies or customers through the app."
                },
                {
                    "name": "Salary & Payroll Export",
                    "detail": "Salaries are not paid out directly through RafApp. Instead, approved timesheets, leave logs, and overtime calculations are compiled and exported as payroll files (CSV or XML) to be uploaded into bank portals or separate accounting tools."
                }
            ]
        },
        {
            "num": "13",
            "title": "Microsoft Business Central Integration",
            "desc": "Synchronizing project tracking, time logs, and material purchases with enterprise ERP systems.",
            "actions": [
                {
                    "name": "ERP Job Cards Sync",
                    "detail": "RafApp integrates with Microsoft Dynamics 365 Business Central via custom REST API connectors and Microsoft Dataverse."
                },
                {
                    "name": "Azure AD OAuth 2.0 Setup",
                    "detail": "Administrators configure the link in the Integrations Panel by authenticating via Microsoft Azure Active Directory (OAuth 2.0) and mapping RafApp entities (Projects, Users, Materials) to Business Central Job Cards, Resource journals, and Inventory tables."
                },
                {
                    "name": "Automated Material Ledger Postings",
                    "detail": "Once set up, approved timesheets automatically register as resource entries in Business Central, and material lists trigger draft purchase invoice orders."
                }
            ]
        },
        {
            "num": "14",
            "title": "Third-Party Application & API Integrations",
            "desc": "Configuring connections with external web services, database endpoints, and wholesaler books.",
            "actions": [
                {
                    "name": "API Key Credentials",
                    "detail": "Superadmins generate API access keys with specific permissions to query catalogs, fetch time logs, or push project updates from external custom applications."
                },
                {
                    "name": "Webhooks Event Triggers",
                    "detail": "Configure webhooks to notify external services (e.g. Slack, MS Teams, custom endpoints) in real-time when tasks are finished, tool faults are reported, or material orders are signed off."
                },
                {
                    "name": "Wholesaler Price Book Sync",
                    "detail": "Background cron tasks connect to wholesale suppliers (Ískraft, Reykjafell, and Johan Ronning) via daily XML/JSON feeds to fetch and update spec sheets and price books."
                }
            ]
        },
        {
            "num": "15",
            "title": "Mobile Experience & Google Play Store Publishing",
            "desc": "Operating and deploying RafApp's mobile web client across user devices.",
            "actions": [
                {
                    "name": "Progressive Web App (PWA) Operation",
                    "detail": "RafApp is a Progressive Web App (PWA) with home screen installation, local resource caching, offline timesheet logging, and modern page loads."
                },
                {
                    "name": "Google Play Store Publishing",
                    "detail": "The PWA can be packaged for Google Play Store distribution using Trusted Web Activity (TWA) wrappers (such as Bubblewrap). Superadmins compile the app bundle (.aab), set up Digital Asset Links, and upload to the Google Play Console for store indexing."
                }
            ]
        }
    ]

    for ch in modules:
        # Add Chapter Heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(8)
        run = h.add_run(f"Module {ch['num']}: {ch['title']}")
        set_font(run, size=15, bold=True, color=RGBColor(30, 41, 59)) # Slate 800

        # Description
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(12)
        run = desc_p.add_run(ch['desc'])
        set_font(run, size=11, color=RGBColor(75, 85, 99)) # Gray 600

        # Add Actions
        for act in ch['actions']:
            sh = doc.add_paragraph()
            sh.paragraph_format.left_indent = Inches(0.2)
            sh.paragraph_format.space_before = Pt(6)
            sh.paragraph_format.space_after = Pt(3)
            run = sh.add_run(f"✦ {act['name']}")
            set_font(run, size=12, bold=True, color=RGBColor(79, 70, 229)) # Indigo 600

            detail_p = doc.add_paragraph()
            detail_p.paragraph_format.left_indent = Inches(0.4)
            detail_p.paragraph_format.space_after = Pt(8)
            run = detail_p.add_run(act['detail'])
            set_font(run, size=10.5, color=RGBColor(30, 41, 59))

    doc.save("C:/Users/mario/Desktop/RafApp_User_Guide.docx")
    print("Successfully generated Word Document at C:/Users/mario/Desktop/RafApp_User_Guide.docx")

if __name__ == "__main__":
    create_professional_guide()
