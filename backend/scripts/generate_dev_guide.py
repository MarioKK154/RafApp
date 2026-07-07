import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_deep_developer_guide():
    doc = docx.Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style helper
    def set_font(run, name="Consolas", size=10, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def set_font_body(run, name="Segoe UI", size=10.5, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(10)
    run = title_p.add_run("RAFAPP ENTERPRISE ARCHITECTURE SPECIFICATION")
    set_font_body(run, size=24, bold=True, color=RGBColor(9, 79, 164))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(200)
    run = subtitle_p.add_run("A Deep-Dive Technical Blueprint covering Schemas, Algorithmic Scoring,\nEncoding Healing, PWA Cache Busting, and Sandbox Isolation")
    set_font_body(run, size=13, italic=True, color=RGBColor(107, 114, 128))

    doc.add_page_break()

    # Introduction
    intro_h = doc.add_paragraph()
    intro_h.paragraph_format.space_before = Pt(12)
    intro_h.paragraph_format.space_after = Pt(8)
    run = intro_h.add_run("System Overview & Codebase Architecture")
    set_font_body(run, size=16, bold=True, color=RGBColor(30, 41, 59))

    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(18)
    run = intro_p.add_run(
        "RafApp is a distributed SaaS application optimized for micro-businesses in the electrical sector. "
        "The architecture is designed for multi-tenancy, high-performance database querying, and robust "
        "data ingestion validation. This document outlines the codebase's deep technical logic, including "
        "SQL queries, database schemas, and client-side update lifecycles."
    )
    set_font_body(run, size=11, color=RGBColor(55, 65, 81))

    # Modules
    modules = [
        {
            "num": "1",
            "title": "SQLAlchemy Schema & Tenant Isolation",
            "desc": "How RafApp separates customer databases logically and configures feature modularity.",
            "points": [
                {
                    "name": "Database Table Definitions (SQLAlchemy Model Snippet)",
                    "detail": (
                        "The primary models use SQLAlchemy declarative definitions. For example, the Tenant model:\n\n"
                        "class Tenant(Base):\n"
                        "    __tablename__ = 'tenants'\n"
                        "    id = Column(Integer, primary_key=True, index=True)\n"
                        "    name = Column(String, nullable=False)\n"
                        "    enabled_features = Column(ARRAY(String), default=list)\n"
                        "    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))\n\n"
                        "Transactions contain 'tenant_id = Column(Integer, ForeignKey(\"tenants.id\"), nullable=False)' "
                        "and are automatically isolated in CRUD endpoints via route-level filters: 'query.filter(m.tenant_id == active_tenant_id)'."
                    )
                },
                {
                    "name": "Multi-Database Engine Configurations",
                    "detail": (
                        "SQLAlchemy connection pool routing is handled in app/database.py using role mappings:\n"
                        "engines_by_role = {\n"
                        "    'primary': create_engine(DATABASE_URL),\n"
                        "    'registry': create_engine(DATABASE_URL_REGISTRY),\n"
                        "    'shop': create_engine(DATABASE_URL_SHOP),\n"
                        "    'reference': create_engine(DATABASE_URL_REFERENCE)\n"
                        "}\n"
                        "This permits database scale-out (e.g. putting the shop catalog and logs on dedicated databases)."
                    )
                }
            ]
        },
        {
            "num": "2",
            "title": "Token Security, Impersonation & Security Compliance",
            "desc": "Authentication workflows, administration security scopes, and compliance notifications.",
            "points": [
                {
                    "name": "Impersonation Logic & JWT Payload Injection",
                    "detail": (
                        "To perform troubleshooting, superadmins trigger a token swap at '/api/admin/impersonate/{user_id}'. "
                        "The system returns a JWT payload containing:\n"
                        "{\n"
                        "  'sub': 'user_id_to_impersonate',\n"
                        "  'tenant_id': active_tenant_id,\n"
                        "  'role': target_user_role,\n"
                        "  'is_impersonating': true,\n"
                        "  'admin_actor_id': superadmin_user_id\n"
                        "}"
                    )
                },
                {
                    "name": "Audit Logs Injection",
                    "detail": (
                        "When the impersonation session starts, the backend executes an automatic insert:\n"
                        "db.add(Notification(\n"
                        "    user_id=target_user_id,\n"
                        "    title='Security Alert: Session Impersonation',\n"
                        "    message='A system administrator has entered your account to perform updates.',\n"
                        "    level='warning'\n"
                        "))\n"
                        "This ensures strict GDPR/SaaS compliance: users are instantly aware of support intrusions."
                    )
                }
            ]
        },
        {
            "num": "3",
            "title": "Algorithmic Search Relevance Engine",
            "desc": "How search queries are parsed into tokens, expanded via synonyms, and scored dynamically using SQL.",
            "points": [
                {
                    "name": "Regex Token Generation (inventory_search.py)",
                    "detail": (
                        "Regex patterns extract and swap electrical variables. For instance, swapping digit-bounded separators:\n"
                        "re.sub(r'(?<=\\d)([gGxX])(?=\\d)', lambda m: 'x' if m.group(1).lower() == 'g' else 'g', s)\n"
                        "This expands '3g2,5' into primary patterns: ['3g2.5', '3g2,5', '3x2.5', '3x2,5']. Synonyms map to secondary arrays."
                    )
                },
                {
                    "name": "SQL CASE WHEN Relevance Calculation",
                    "detail": (
                        "The SQLAlchemy ORM compiles a custom CASE WHEN structure to rank products. The compiled SQL query looks like:\n\n"
                        "SELECT name, subcategory, (\n"
                        "  (CASE WHEN name ILIKE '%nym%' OR name_en ILIKE '%nym%' THEN 100 ELSE 0 END) +\n"
                        "  (CASE WHEN name ILIKE '%nym-j%' OR name_en ILIKE '%nym-j%' THEN 50 ELSE 0 END) +\n"
                        "  (CASE WHEN name ILIKE '%mmj%' OR name ILIKE '%exq%' THEN 10 ELSE 0 END) +\n"
                        "  (CASE WHEN description ILIKE '%nym%' OR ronning_sku ILIKE '%nym%' THEN 5 ELSE 0 END)\n"
                        ") AS relevance_score\n"
                        "FROM inventory_items\n"
                        "WHERE name ILIKE '%nym%' OR name ILIKE '%mmj%' OR name ILIKE '%exq%'\n"
                        "ORDER BY relevance_score DESC, name ASC;"
                    )
                }
            ]
        },
        {
            "num": "4",
            "title": "Encoding Healing & Ingestion Pipeline",
            "desc": "How data is sanitized and restored when parsed from corrupted/double-encoded Excel tables.",
            "points": [
                {
                    "name": "Double-Encoding Cleaner Logic",
                    "detail": (
                        "When database exports are subjected to double serialization, Icelandic characters become corrupted. The recovery logic is:\n\n"
                        "def restore_string(s: str) -> str:\n"
                        "    # Check if the string contains double UTF-8 signature bytes (like 0xc2 or 0xc3)\n"
                        "    if any(ord(c) in (0xc2, 0xc3) for c in s):\n"
                        "        try:\n"
                        "            # Re-encode to original bytes and decode back as UTF-8\n"
                        "            return s.encode('latin-1').decode('utf-8')\n"
                        "        except (UnicodeEncodeError, UnicodeDecodeError):\n"
                        "            return s\n"
                        "    return s\n"
                        "This heals strings like 'Ryfrtt' back to 'Ryðfrítt' dynamically."
                    )
                },
                {
                    "name": "Excel Multi-Sheet Parsing",
                    "detail": (
                        "The import pipeline opens the Excel workbook, selects target worksheets ('Cables', 'Cable trays and ladders', 'Pipes'), "
                        "and processes records. It configures English names as primary stable filter keys inside the 'category' and 'subcategory' fields, "
                        "storing localized Icelandic strings in '_en' counterpart columns. This prevents category collision when names repeat."
                    )
                }
            ]
        },
        {
            "num": "5",
            "title": "Cache Eviction & Hot-Reload Client Lifecycles",
            "desc": "How frontend clients wipe browser caches and reload when code updates are deployed.",
            "points": [
                {
                    "name": "Client-Side Cache Eviction Script (Javascript)",
                    "detail": (
                        "Vite compiles a build timestamp. On mount, App.jsx compares this to the server: '/api/system/version'. If mismatched, it executes:\n\n"
                        "if (navigator.serviceWorker) {\n"
                        "    navigator.serviceWorker.getRegistrations().then(regs => {\n"
                        "        for (let reg of regs) reg.unregister();\n"
                        "    });\n"
                        "}\n"
                        "if (window.caches) {\n"
                        "    caches.keys().then(keys => {\n"
                        "        keys.forEach(k => caches.delete(k));\n"
                        "    });\n"
                        "}\n"
                        "localStorage.setItem('app_version', serverVersion);\n"
                        "window.location.reload(true); // force browser to reload static bundle files"
                    )
                }
            ]
        },
        {
            "num": "6",
            "title": "Testing Sandbox & Database Isolation",
            "desc": "How unit tests are isolated from Postgres tables during pytest runs.",
            "points": [
                {
                    "name": "Force SQLite DB Mocking (conftest.py)",
                    "detail": (
                        "To protect production PostgreSQL data, conftest.py intercepts db configs before tests boot up:\n\n"
                        "import os\n"
                        "os.environ['DATABASE_URL'] = 'sqlite:///./test.db'\n"
                        "os.environ['DATABASE_URL_REGISTRY'] = 'sqlite:///./test.db'\n"
                        "os.environ['DATABASE_URL_SHOP'] = 'sqlite:///./test.db'\n\n"
                        "# SQLAlchemy then creates local SQLite tables. Every test runs within a transaction block:\n"
                        "@pytest.fixture(scope='function')\n"
                        "def db():\n"
                        "    connection = engine.connect()\n"
                        "    transaction = connection.begin()\n"
                        "    session = Session(bind=connection)\n"
                        "    yield session\n"
                        "    session.close()\n"
                        "    transaction.rollback()\n"
                        "    connection.close()"
                    )
                }
            ]
        }
    ]

    for ch in modules:
        # Add Chapter Heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(8)
        run = h.add_run(f"Section {ch['num']}: {ch['title']}")
        set_font_body(run, size=15, bold=True, color=RGBColor(30, 41, 59)) # Slate 800

        # Description
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(12)
        run = desc_p.add_run(ch['desc'])
        set_font_body(run, size=11, color=RGBColor(75, 85, 99)) # Gray 600

        # Add Points
        for pt in ch['points']:
            sh = doc.add_paragraph()
            sh.paragraph_format.left_indent = Inches(0.2)
            sh.paragraph_format.space_before = Pt(6)
            sh.paragraph_format.space_after = Pt(3)
            run = sh.add_run(f"✦ {pt['name']}")
            set_font_body(run, size=12, bold=True, color=RGBColor(9, 79, 164)) # Deep Blue 600

            # For Code blocks, we use Consolas font
            detail_p = doc.add_paragraph()
            detail_p.paragraph_format.left_indent = Inches(0.4)
            detail_p.paragraph_format.space_after = Pt(8)
            
            # Simple check if there's code-like content
            lines = pt['detail'].split('\n')
            for idx, line in enumerate(lines):
                if any(k in line for k in ['def ', 'class ', 'import ', 'SELECT ', 'ORDER BY', '{', '}', 'if (', 'const ']):
                    run_line = detail_p.add_run(line + '\n')
                    set_font(run_line, size=9.5, color=RGBColor(5, 120, 80)) # Greenish console look
                else:
                    run_line = detail_p.add_run(line + '\n')
                    set_font_body(run_line, size=10.5, color=RGBColor(30, 41, 59))

    doc.save("C:/Users/mario/Desktop/RafApp_Developer_Guide.docx")
    print("Successfully generated Deep Word Document at C:/Users/mario/Desktop/RafApp_Developer_Guide.docx")

if __name__ == "__main__":
    create_deep_developer_guide()
